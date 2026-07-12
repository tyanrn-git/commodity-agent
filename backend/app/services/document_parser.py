import email
import hashlib
import re
from email import policy
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".eml"}
ALLOWED_UPLOAD_MIME = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "message/rfc822",
    "application/octet-stream",
}


def compute_content_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def extract_text_from_bytes(*, filename: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension == ".pdf":
        return _extract_pdf_text(content)
    if extension == ".docx":
        return _extract_docx_text(content)
    if extension == ".xlsx":
        return _extract_xlsx_text(content)
    if extension == ".eml":
        return _extract_eml_text(content)
    raise ValueError(f"Unsupported file type: {extension}")


def _extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
    result = "\n".join(parts).strip()
    if not result:
        raise ValueError("PDF contains no extractable text (OCR not supported in MVP)")
    return result


def _extract_docx_text(content: bytes) -> str:
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    result = "\n".join(parts).strip()
    if not result:
        raise ValueError("DOCX contains no extractable text")
    return result


def _extract_xlsx_text(content: bytes) -> str:
    workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in workbook.worksheets:
        parts.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                parts.append(" | ".join(values))
    result = "\n".join(parts).strip()
    if not result:
        raise ValueError("XLSX contains no extractable text")
    return result


def _extract_eml_text(content: bytes) -> str:
    message = email.message_from_bytes(content, policy=policy.default)
    parts: list[str] = []
    subject = message.get("subject")
    if subject:
        parts.append(f"Subject: {subject}")
    if message.is_multipart():
        for part in message.walk():
            if part.get_content_type() == "text/plain":
                payload = part.get_content()
                if isinstance(payload, str) and payload.strip():
                    parts.append(payload.strip())
    else:
        payload = message.get_content()
        if isinstance(payload, str) and payload.strip():
            parts.append(payload.strip())
    result = "\n".join(parts).strip()
    if not result:
        raise ValueError("EML contains no extractable text")
    return result


def fetch_public_url_text(url: str, *, timeout: float = 15.0) -> tuple[str, bytes]:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise ValueError("Only http/https URLs are supported")
    if not parsed.netloc:
        raise ValueError("Invalid URL")

    with httpx.Client(timeout=timeout, follow_redirects=True) as client:
        response = client.get(url, headers={"User-Agent": "CommodityAgent/1.0"})
        response.raise_for_status()
        content_type = response.headers.get("content-type", "")
        if "html" not in content_type.lower():
            raise ValueError("URL import supports only public HTML pages in MVP")
        html = response.text
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = re.sub(r"\n{3,}", "\n\n", soup.get_text("\n", strip=True))
        if not text.strip():
            raise ValueError("No text extracted from URL")
        encoded = text.encode("utf-8")
        return text, encoded
